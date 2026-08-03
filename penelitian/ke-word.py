#!/usr/bin/env python3
"""
proposal.md → proposal.docx dengan format Panduan Penulisan Disertasi FK USU.

Dijalankan dari akar proyek:

    python3 penelitian/ke-word.py

Yang DIKERJAKAN skrip ini: tata huruf, spasi, batas tepi, gaya judul bab dan
sub bab, tabel bergaris atas-bawah saja, gambar tersemat beserta keterangannya,
dan penomoran halaman.

Yang TIDAK dikerjakan, karena bentuknya ditentukan fakultas dan tidak pantas
dikarang sendiri: halaman sampul, lembar pengesahan, pernyataan orisinalitas,
abstrak, daftar isi, daftar tabel, daftar gambar, dan daftar singkatan. Semua
itu ditambahkan sendiri memakai templat resmi FK USU.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AKAR = Path(__file__).resolve().parent
SUMBER = AKAR / 'proposal.md'
HASIL = AKAR / 'proposal.docx'
GAMBAR = AKAR / 'gambar'

# Panduan FK USU bagian 3: Times New Roman 12, judul tabel/gambar 10.
FONT = 'Times New Roman'
UKURAN = Pt(12)
UKURAN_KECIL = Pt(10)
# "Mulai ketikan ke-6 dari tepi kiri" — enam ketukan Times New Roman 12.
INDEN = Cm(1.25)
MERAH = RGBColor(0xC0, 0x00, 0x00)

# Blok "Catatan pengerjaan" pada naskah diganti gambar yang sesungguhnya.
# Kunci = potongan nama berkas yang disebut di dalam catatan itu.
GAMBAR_SETELAH_SUBBAB = {
    '3.12 Alur Penelitian': ['gambar-3-1', 'gambar-3-2'],
}

KETERANGAN_GAMBAR = {
    'gambar-2-1': ('Gambar 2.1', 'Kerangka teori'),
    'gambar-2-2': ('Gambar 2.2', 'Kerangka konsep'),
    'gambar-3-1': ('Gambar 3.1', 'Alur rekrutmen dan persetujuan'),
    'gambar-3-2': ('Gambar 3.2', 'Alur pemantauan dan analisis'),
}


def atur_gaya(doc: Document) -> None:
    n = doc.styles['Normal']
    n.font.name = FONT
    n.font.size = UKURAN
    # Nama font harus ditulis juga untuk aksara Asia Timur, kalau tidak Word
    # diam-diam memakai font lain pada sebagian karakter.
    n.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    p = n.paragraph_format
    p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.space_after = Pt(0)
    p.space_before = Pt(0)


def atur_halaman(sec) -> None:
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)   # Panduan menebalkan yang ini.
    sec.right_margin = Cm(3)


def nomor_halaman_kanan_atas(sec) -> None:
    """Angka Arab di kanan atas, sesuai penomoran bagian utama."""
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    r.font.name = FONT
    r.font.size = UKURAN
    for isi, atribut in (('begin', None), (None, 'PAGE'), ('end', None)):
        el = OxmlElement('w:fldChar') if isi else OxmlElement('w:instrText')
        if isi:
            el.set(qn('w:fldCharType'), isi)
        else:
            el.set(qn('xml:space'), 'preserve')
            el.text = atribut
        r._r.append(el)


def garis_atas_bawah(tabel) -> None:
    """
    Tabel tanpa garis pemisah vertikal; garis batas hanya di atas dan di bawah.

    Panduan bagian 5 butir 3. Word tidak punya saklar untuk itu, jadi tepi tiap
    sel disetel satu per satu.
    """
    baris_akhir = len(tabel.rows) - 1
    for i, baris in enumerate(tabel.rows):
        for sel in baris.cells:
            tcPr = sel._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for sisi in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{sisi}')
                tampil = (sisi == 'top' and i == 0) or (sisi == 'bottom' and i == baris_akhir)
                # Garis bawah baris kepala juga dipertahankan: tanpa itu
                # kepala tabel menyatu dengan isinya dan sulit dibaca.
                if sisi == 'bottom' and i == 0:
                    tampil = True
                el.set(qn('w:val'), 'single' if tampil else 'none')
                el.set(qn('w:sz'), '8')
                el.set(qn('w:color'), '000000')
                borders.append(el)
            tcPr.append(borders)


def tulis_inline(par, teks: str, merah: bool = False) -> None:
    """
    Menerjemahkan **tebal**, *miring*, dan `berkas` menjadi runs.

    Backtick dijadikan miring, bukan huruf mesin ketik: yang ada di dalamnya
    adalah nama berkas, dan panduan meminta kata asing dicetak miring.
    """
    for bagian in re.split(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)', teks):
        if not bagian:
            continue
        r = par.add_run()
        if bagian.startswith('**') and bagian.endswith('**'):
            r.text, r.bold = bagian[2:-2], True
        elif bagian.startswith('*') and bagian.endswith('*'):
            r.text, r.italic = bagian[1:-1], True
        elif bagian.startswith('`') and bagian.endswith('`'):
            r.text, r.italic = bagian[1:-1], True
        else:
            r.text = bagian
        r.font.name = FONT
        r.font.size = UKURAN
        if merah:
            r.font.color.rgb = MERAH


def paragraf(doc, teks, *, inden=True, rata=WD_ALIGN_PARAGRAPH.JUSTIFY, merah=False):
    p = doc.add_paragraph()
    p.alignment = rata
    if inden:
        p.paragraph_format.first_line_indent = INDEN
    tulis_inline(p, teks, merah=merah)
    return p


def judul_bab(doc, teks: str) -> None:
    """
    Judul bab: kapital semua, simetris, tebal, di halaman baru.

    "BAB I PENDAHULUAN" dipecah dua baris sesuai kebiasaan naskah FK USU.
    """
    doc.add_page_break()
    cocok = re.match(r'^(BAB [IVX]+)\s+(.*)$', teks)
    baris = [cocok.group(1), cocok.group(2)] if cocok else [teks]
    for b in baris:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(b.upper())
        r.bold = True
        r.font.name = FONT
        r.font.size = UKURAN
    doc.add_paragraph()


def judul_sub(doc, teks: str, anak: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(teks)
    r.bold = True
    r.font.name = FONT
    r.font.size = UKURAN
    if anak:
        r.italic = True


def sisipkan_gambar(doc, kunci: str) -> bool:
    berkas = GAMBAR / f'{kunci}.png'
    if not berkas.exists():
        return False

    nomor, judul = KETERANGAN_GAMBAR[kunci]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    # Lebar bidang ketik A4 dengan tepi 4 cm dan 3 cm adalah 14 cm.
    p.add_run().add_picture(str(berkas), width=Cm(14))

    # Judul gambar di BAWAH gambar, tanpa titik, sentence case, 1 spasi, 10 pt.
    k = doc.add_paragraph()
    k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    k.paragraph_format.space_after = Pt(12)
    r = k.add_run(f'{nomor} {judul}')
    r.font.name = FONT
    r.font.size = UKURAN_KECIL
    return True


def sisipkan_tabel(doc, baris_md, nomor_judul):
    kolom = [[s.strip() for s in b.strip().strip('|').split('|')] for b in baris_md]
    isi = [k for k in kolom if not all(re.fullmatch(r':?-{2,}:?', s) for s in k)]

    # Judul tabel di ATAS tabel, diakhiri titik, 1 spasi, 10 pt.
    j = doc.add_paragraph()
    j.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = j.add_run(nomor_judul)
    r.font.name = FONT
    r.font.size = UKURAN_KECIL

    t = doc.add_table(rows=len(isi), cols=len(isi[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, brs in enumerate(isi):
        for k, sel in enumerate(brs):
            c = t.cell(i, k)
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            tulis_inline(p, sel)
            for run in p.runs:
                run.font.size = UKURAN_KECIL
                if i == 0:
                    run.bold = True
    garis_atas_bawah(t)
    doc.add_paragraph()
    return t


def main() -> int:
    if not SUMBER.exists():
        print(f'Tidak menemukan {SUMBER}', file=sys.stderr)
        return 1

    doc = Document()
    atur_gaya(doc)
    atur_halaman(doc.sections[0])

    baris = SUMBER.read_text(encoding='utf-8').split('\n')
    i = 0
    nomor_tabel = 0
    halaman_utama_dimulai = False
    kutipan_tertunda: list[str] = []
    subbab_kini = ''
    gambar_tertunda: list[str] = []
    # Daftar pustaka dan lampiran punya aturan spasi sendiri.
    di_daftar_pustaka = False

    def buang_kutipan():
        """Blok `>` yang bukan penanda gambar ditulis merah sebagai catatan kerja."""
        nonlocal kutipan_tertunda
        if not kutipan_tertunda:
            return
        gabung = ' '.join(kutipan_tertunda)

        for kunci in KETERANGAN_GAMBAR:
            if kunci in gabung:
                sisipkan_gambar(doc, kunci)
                kutipan_tertunda = []
                return

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = INDEN
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Catatan kerja dan hipotesis sama-sama blok kutipan di sumbernya.
        # Yang jelas-jelas catatan diberi warna merah supaya tidak mungkin
        # ikut terkirim ke komite etik tanpa disadari.
        catatan = 'Catatan' in gabung or 'catatan kerja' in gabung
        tulis_inline(p, gabung, merah=catatan)
        doc.add_paragraph()
        kutipan_tertunda = []

    while i < len(baris):
        b = baris[i]
        s = b.strip()

        if s.startswith('>'):
            kutipan_tertunda.append(s.lstrip('> ').rstrip())
            i += 1
            continue
        buang_kutipan()

        if not s or s == '---':
            i += 1
            continue

        if s.startswith('# '):
            teks = s[2:].strip()
            if not halaman_utama_dimulai and not teks.startswith('BAB'):
                # Judul penelitian pada halaman pertama.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(teks.upper())
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(14)
                doc.add_paragraph()
            else:
                for kunci in gambar_tertunda:
                    sisipkan_gambar(doc, kunci)
                gambar_tertunda = []
                di_daftar_pustaka = teks.strip().upper().startswith('DAFTAR PUSTAKA')
                if not halaman_utama_dimulai:
                    # Bagian utama dimulai: sesi baru supaya penomoran Arab
                    # tidak ikut menomori halaman judul.
                    sec = doc.add_section(WD_SECTION.NEW_PAGE)
                    atur_halaman(sec)
                    sec.header.is_linked_to_previous = False
                    nomor_halaman_kanan_atas(sec)
                    halaman_utama_dimulai = True
                judul_bab(doc, teks)
            i += 1
            continue

        if s.startswith('### '):
            judul_sub(doc, s[4:].strip(), anak=True)
            i += 1
            continue

        if s.startswith('## '):
            # Gambar milik sub bab sebelumnya disisipkan sebelum pindah.
            for kunci in gambar_tertunda:
                sisipkan_gambar(doc, kunci)
            gambar_tertunda = []

            subbab_kini = s[3:].strip()
            judul_sub(doc, subbab_kini)
            gambar_tertunda = list(GAMBAR_SETELAH_SUBBAB.get(subbab_kini, []))
            i += 1
            continue

        if s.startswith('|'):
            blok = []
            while i < len(baris) and baris[i].strip().startswith('|'):
                blok.append(baris[i])
                i += 1
            nomor_tabel += 1
            # Judul diambil dari sub bab tempat tabelnya berada, supaya tidak
            # perlu diperbarui manual setiap naskah berubah.
            nama = re.sub(r'^Lampiran \d+\.\s*', '', subbab_kini).strip()
            nama = nama[0].upper() + nama[1:] if nama else 'Tabel'
            sisipkan_tabel(doc, blok, f'Tabel {nomor_tabel}. {nama}.')
            continue

        cocok_daftar = re.match(r'^(\d+)\.\s+(.*)$', s)
        if cocok_daftar:
            # Baris lanjutan sebuah butir menjorok di sumbernya.
            isi = [cocok_daftar.group(2)]
            i += 1
            while i < len(baris) and baris[i].startswith('   ') and baris[i].strip():
                isi.append(baris[i].strip())
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = INDEN
            p.paragraph_format.first_line_indent = -INDEN
            tulis_inline(p, f'{cocok_daftar.group(1)}. ' + ' '.join(isi))
            continue

        # Paragraf biasa: kumpulkan sampai baris kosong.
        isi = [s]
        i += 1
        while i < len(baris) and baris[i].strip() and not re.match(
            r'^\s*(#|>|\||\d+\.\s|---)', baris[i]
        ):
            isi.append(baris[i].strip())
            i += 1

        if di_daftar_pustaka:
            # Panduan bagian 3: daftar pustaka memakai satu spasi. Barisnya
            # dibuat menggantung supaya nama pengarang menonjol ke kiri.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.left_indent = INDEN
            p.paragraph_format.first_line_indent = -INDEN
            p.paragraph_format.space_after = Pt(6)
            tulis_inline(p, ' '.join(isi))
        else:
            paragraf(doc, ' '.join(isi))

    buang_kutipan()
    for kunci in gambar_tertunda:
        sisipkan_gambar(doc, kunci)
    doc.save(HASIL)
    print(f'Tersimpan: {HASIL}')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
